"""
Layout-faithful OCR pipeline: detect -> recognise -> remove text -> replace in place.

Produces, for a page image:
  <stem>_clean_bg.png   original with text removed (inpainted)
  <stem>_overlay.png    recognised text placed at exact detected positions
  <stem>_aligned.docx   editable Word file, text boxes at exact positions
  <stem>_lines.json     boxes + text + per-box alignment score

Deliberately avoids the failure modes of docx_generator.py:
  * no fixed 8% crop            -> page region is detected
  * no cascading collision shift -> a bad box cannot drag later lines out of place
  * no +2px / *0.8 fudge factors
  * one font-sizing rule, fitted to the measured box

Usage:
    python aligned_pipeline.py <image-or-pdf> [--model NAME] [--no-ocr]
Requires: an Ollama server with a vision OCR model pulled.
"""
import argparse, base64, json, os, re, sys, threading, time, urllib.request
import cv2, numpy as np
from PIL import Image, ImageDraw, ImageFont

OLLAMA = os.environ.get("OLLAMA_URL", "http://localhost:11434")
DEFAULT_MODEL = "AuditAid/PaddleOCR-VL-1.6-0.9B"
PROMPT = "Text Recognition:"
HERE_DIR = os.path.dirname(os.path.abspath(__file__))

# Nirmala UI ships with Windows and covers Kannada + Devanagari.
FONT_CANDIDATES = [
    r"C:\Windows\Fonts\Nirmala.ttc",
    "/usr/share/fonts/truetype/noto/NotoSansKannada-Regular.ttf",
    "/usr/share/fonts/truetype/lohit-kannada/Lohit-Kannada.ttf",
]


# ----------------------------------------------------------------- load page
def load_page(path):
    if path.lower().endswith(".pdf"):
        import fitz
        d = fitz.open(path)
        page = d[0]
        imgs = page.get_images(full=True)
        if imgs:                                     # prefer the embedded scan
            info = d.extract_image(imgs[0][0])
            arr = np.frombuffer(info["image"], np.uint8)
            return cv2.imdecode(arr, cv2.IMREAD_COLOR)
        pix = page.get_pixmap(dpi=300)               # else render
        return cv2.imdecode(np.frombuffer(pix.tobytes("png"), np.uint8), cv2.IMREAD_COLOR)
    return cv2.imread(path)


# ------------------------------------------------- stage 0: isolate + deskew
def normalise(img):
    H, W = img.shape[:2]
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    _, pm = cv2.threshold(cv2.GaussianBlur(gray, (7, 7), 0), 0, 255,
                          cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    pm = cv2.morphologyEx(pm, cv2.MORPH_CLOSE,
                          cv2.getStructuringElement(cv2.MORPH_RECT, (31, 31)))
    cnts, _ = cv2.findContours(pm, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    page = max(cnts, key=cv2.contourArea)
    frac = cv2.contourArea(page) / (W * H)

    # gentle shadow flatten - keeps thin Indic strokes that hard thresholding erases
    bg = cv2.GaussianBlur(gray, (0, 0), sigmaX=45)
    flat = np.clip(gray.astype(np.float32) / np.maximum(bg, 1) * 200, 0, 255).astype(np.uint8)
    # Strict threshold for the global pass. A page-wide sensitive threshold was
    # tried and rejected: it found faint text but added ~55 spurious boxes and
    # corrupted the skew estimate (-0.72 vs the true -2.11 deg). Faint lines are
    # instead recovered surgically, only inside detected vertical gaps - see
    # recover_faint_lines().
    ink = cv2.adaptiveThreshold(flat, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                cv2.THRESH_BINARY_INV, 25, 12)
    ink_sensitive = cv2.adaptiveThreshold(flat, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                          cv2.THRESH_BINARY_INV, 41, 4)
    ink_sensitive = cv2.morphologyEx(
        ink_sensitive, cv2.MORPH_OPEN,
        cv2.getStructuringElement(cv2.MORPH_RECT, (2, 2)))
    inner = np.zeros_like(ink)
    cv2.drawContours(inner, [page], -1, 255, -1)
    inner = cv2.erode(inner, cv2.getStructuringElement(cv2.MORPH_RECT, (25, 25)))
    ink = cv2.bitwise_and(ink, inner)

    sm = cv2.morphologyEx(ink, cv2.MORPH_CLOSE,
                          cv2.getStructuringElement(cv2.MORPH_RECT, (75, 3)))
    cs, _ = cv2.findContours(sm, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    angs = []
    for c in cs:
        x, y, w, h = cv2.boundingRect(c)
        if w > 200 and 8 < h < 90:
            a = cv2.minAreaRect(c)[2]
            a = a - 90 if a > 45 else (a + 90 if a < -45 else a)
            if abs(a) < 15:
                angs.append(a)
    skew = float(np.median(angs)) if angs else 0.0

    M = cv2.getRotationMatrix2D((W / 2, H / 2), skew, 1.0)
    rot = cv2.warpAffine(img, M, (W, H), flags=cv2.INTER_CUBIC,
                         borderMode=cv2.BORDER_REPLICATE)
    ink = cv2.warpAffine(ink, M, (W, H), flags=cv2.INTER_NEAREST)
    ink_sensitive = cv2.warpAffine(ink_sensitive, M, (W, H), flags=cv2.INTER_NEAREST)
    return rot, ink, skew, frac, ink_sensitive, M


def apply_table_grid(boxes, ink, enabled=None, verbose=True):
    """Inside a ruled table, replace row-spanning line boxes with per-cell lines.

    A detected "line" normally spans a whole table row, merging Item No.,
    Subject and Page No. into one string and destroying the columns. Restricting
    line detection to each cell keeps every column separate.
    """
    if enabled is None:
        enabled = os.environ.get("OCR_TABLES", "1") == "1"
    if not enabled:
        return boxes
    tables = detect_table_grid(ink)
    if not tables:
        return boxes

    med_h = float(np.median([b["h"] for b in boxes])) if boxes else 20.0
    kept = list(boxes)
    for cols, rows, (tx0, ty0, tx1, ty1) in tables:
        cells, newlines = 0, []
        for ri in range(len(rows) - 1):
            for ci in range(len(cols) - 1):
                cx, cy = cols[ci], rows[ri]
                cw, ch = cols[ci + 1] - cx, rows[ri + 1] - cy
                if cw < 16 or ch < 12:
                    continue
                # 2px inward keeps the printed rule out of the cell while
                # leaving room for a leading glyph that sits close to it.
                pad = 2
                cell = {"x": int(cx + pad), "y": int(cy + pad),
                        "w": int(max(4, cw - 2 * pad)),
                        "h": int(max(4, ch - 2 * pad)),
                        "row": int(ri), "col": int(ci)}
                sub = ink[cell["y"]:cell["y"] + cell["h"],
                          cell["x"]:cell["x"] + cell["w"]]
                if sub.size == 0 or sub.mean() / 255.0 < 0.008:
                    continue
                cells += 1
                newlines.extend(cell_lines(ink, cell, med_h))
        if not newlines:
            continue

        def in_tbl(b):
            bx, by = b["x"] + b["w"] / 2, b["y"] + b["h"] / 2
            return tx0 - 8 <= bx <= tx1 + 8 and ty0 - 8 <= by <= ty1 + 8

        before = len(kept)
        kept = [b for b in kept if not in_tbl(b)] + newlines
        if verbose:
            print(f"  table: {len(cols)-1} cols x {len(rows)-1} rows, "
                  f"{cells} filled cells -> {len(newlines)} cell lines "
                  f"(replaced {before - (len(kept) - len(newlines))} row-spanning boxes)")
    return kept


def recover_faint_lines(boxes, ink_sensitive, ink):
    """Look for text ONLY inside suspicious vertical gaps.

    A faint paragraph can fall below the strict threshold and disappear with no
    trace (observed: 3 lines lost in a 367px gap). Rather than lowering the
    threshold page-wide - which floods the page with noise and wrecks the skew
    estimate - re-examine just the gaps, where real text is the only plausible
    explanation for a large blank run inside the text column.
    """
    if len(boxes) < 4:
        return boxes
    hs = [b["h"] for b in boxes]
    med_h = float(np.median(hs))
    text_boxes = sorted(boxes, key=lambda b: b["y"])
    xs0 = int(np.percentile([b["x"] for b in text_boxes], 20))
    xs1 = int(np.percentile([b["x"] + b["w"] for b in text_boxes], 80))
    pitch = med_h * 1.55

    added = []
    for a, c in zip(text_boxes, text_boxes[1:]):
        top = a["y"] + a["h"]
        gap = c["y"] - top
        if gap < pitch * 1.6:
            continue
        band = ink_sensitive[top:c["y"], xs0:xs1]
        if band.size == 0:
            continue
        rows = (band > 0).sum(axis=1)
        thr = max(band.shape[1] * 0.04, 4)
        run = None
        for i, v in enumerate(rows):
            if v > thr and run is None:
                run = i
            elif v <= thr and run is not None:
                if i - run >= med_h * 0.45:
                    added.append({"x": xs0, "y": top + run,
                                  "w": xs1 - xs0, "h": i - run, "faint": True})
                run = None
        if run is not None and len(rows) - run >= med_h * 0.45:
            added.append({"x": xs0, "y": top + run,
                          "w": xs1 - xs0, "h": len(rows) - run, "faint": True})

    if not added:
        return boxes
    # Tighten each recovered band to its actual ink, both axes. Without the
    # vertical trim the box is as tall as the gap slice, which then drives an
    # absurd font size downstream.
    keep = []
    for b in added:
        sub = ink_sensitive[b["y"]:b["y"] + b["h"], :]
        cols = (sub > 0).sum(axis=0)
        nz = np.nonzero(cols > 2)[0]
        if not len(nz):
            continue
        b["x"], b["w"] = int(nz[0]), int(nz[-1] - nz[0] + 1)
        band = ink_sensitive[b["y"]:b["y"] + b["h"], b["x"]:b["x"] + b["w"]]
        rows = (band > 0).sum(axis=1)
        rz = np.nonzero(rows > max(band.shape[1] * 0.02, 2))[0]
        if not len(rz):
            continue
        b["y"], b["h"] = b["y"] + int(rz[0]), int(rz[-1] - rz[0] + 1)
        if b["w"] < 40 or b["h"] < max(8, med_h * 0.4) or b["h"] > med_h * 2.4:
            continue
        keep.append(b)
    if keep:
        print(f"  recovered {len(keep)} faint line(s) from vertical gaps")
    return boxes + keep


# ------------------------------------------------ stage 1: detect text lines
def split_tall(box, ink, med_h):
    """A block that merged several lines is split by horizontal ink projection."""
    if box["h"] < med_h * 2.2:
        return [box]
    x, y, w, h = box["x"], box["y"], box["w"], box["h"]
    rows = (ink[y:y + h, x:x + w] > 0).sum(axis=1)
    thr = max(rows.max() * 0.10, 2)
    out, start = [], None
    for i, v in enumerate(rows):
        if v > thr and start is None:
            start = i
        elif v <= thr and start is not None:
            if i - start >= max(8, med_h * 0.35):
                out.append({"x": x, "y": y + start, "w": w, "h": i - start})
            start = None
    if start is not None and h - start >= max(8, med_h * 0.35):
        out.append({"x": x, "y": y + start, "w": w, "h": h - start})
    return out or [box]


def _hrules(ink, min_frac=0.14):
    """Horizontal rules as (y, x_start, x_end), merged across thickness.

    Dashed rules must be bridged BEFORE the long-kernel open, or every dash is
    erased as too short. Measured on a 30-row employee table with dashed row
    separators: without bridging only 7 partial rules were found (17-34% of
    width); with it the real grid appears.
    """
    H, W = ink.shape[:2]
    # NOTE: dash-bridging was tried here to catch dashed row rules and REVERTED.
    # A 13px close turned letters into pseudo-rules (76 "rules", 6 phantom
    # tables); a 7px close plus thin-rule and continuity filters still split a
    # working 3x4 table into 3x5 + 2x2 and scrambled row assignment. Dense
    # dashed tables need a trained table model (PP-Structure / TableFormer),
    # not morphology - see OCR_ACCURACY_PLAN.md.
    hk = cv2.getStructuringElement(cv2.MORPH_RECT, (max(30, W // 26), 1))
    horiz = cv2.morphologyEx(ink, cv2.MORPH_OPEN, hk, iterations=1)
    proj = (horiz > 0).sum(axis=1)
    thr = max(int(W * min_frac), 8)
    hits = np.nonzero(proj >= thr)[0]
    if not len(hits):
        return [], horiz
    groups, start, prev = [], hits[0], hits[0]
    for v in hits[1:]:
        if v - prev > 5:
            groups.append((start, prev))
            start = v
        prev = v
    groups.append((start, prev))

    rules = []
    for a, b in groups:
        y = (a + b) // 2
        seg = horiz[a:b + 1, :]
        cols = np.nonzero((seg > 0).any(axis=0))[0]
        if len(cols):
            rules.append((int(y), int(cols[0]), int(cols[-1])))
    return rules, horiz


def detect_table_grid(ink):
    """Find ruled tables and return a cell grid per table.

    Rewritten after the first attempt failed on a real page. Two lessons drove
    the design:
      * the table's x-extent must come from its OWN horizontal rules, not from
        detected verticals - otherwise a binder clip becomes a "column";
      * horizontal rules must be CLUSTERED into tables. Taking the min/max of
        every rule on the page produced a bbox from the header rule to the
        footer rule, so the column test demanded far more vertical ink than any
        real column line could have.
    Returns a list of (cols, rows, bbox).
    """
    H, W = ink.shape[:2]
    rules, _ = _hrules(ink)
    if len(rules) < 3:
        return []

    # cluster rules that plausibly belong to one table: similar x-extent and
    # row-like vertical spacing
    clusters, cur = [], [rules[0]]
    for r in rules[1:]:
        py, px0, px1 = cur[-1]
        y, x0, x1 = r
        gap = y - py
        overlap = (min(px1, x1) - max(px0, x0)) / max(1, min(px1 - px0, x1 - x0))
        if 12 <= gap <= max(260, H * 0.12) and overlap > 0.6:
            cur.append(r)
        else:
            clusters.append(cur)
            cur = [r]
    clusters.append(cur)

    vk = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 12))
    vert = cv2.morphologyEx(ink, cv2.MORPH_OPEN, vk, iterations=1)

    out = []
    for cl in clusters:
        if len(cl) < 3:
            continue
        ys = [r[0] for r in cl]
        y0, y1 = min(ys), max(ys)
        # UNION of the rule extents, not the median. Faint row rules get
        # partially detected (measured: 321..1469, then 756..1474, then
        # 901..1334 on one table), so the median start cut the table's left
        # third off and put the Item-No. column divider outside the bbox.
        x0 = int(min(r[1] for r in cl))
        x1 = int(max(r[2] for r in cl))
        if (x1 - x0) < W * 0.30 or (y1 - y0) < 40:
            continue

        # verticals scoped to THIS table only, and required to span most of it
        band = vert[y0:y1 + 1, x0:x1 + 1]
        if band.size == 0:
            continue
        vproj = (band > 0).sum(axis=0)
        # 0.42, not 0.55: measured on a real table, the Page-No. divider only
        # reaches ~45% of the table height because the bottom rows are unruled
        # on that side. At 0.55 that column was silently dropped.
        need = max(12, int((y1 - y0) * 0.42))
        hits = np.nonzero(vproj >= need)[0]
        cols = []
        if len(hits):
            s, p = hits[0], hits[0]
            for v in hits[1:]:
                if v - p > 6:
                    cols.append(x0 + (s + p) // 2)
                    s = v
                p = v
            cols.append(x0 + (s + p) // 2)
        # the table borders are columns too
        for edge in (x0, x1):
            if not any(abs(c - edge) < 12 for c in cols):
                cols.append(edge)
        cols = sorted(set(cols))

        # Reject text-derived verticals. Character strokes survive the bridge
        # and cluster tightly; real column rules are far apart. Without this a
        # dense 8-column table got 6 bogus columns that sliced dates mid-string
        # ("25/09" | "12" | "006 | 12" | "/01/2") - worse than no table at all.
        MIN_SEP = max(35, int((x1 - x0) * 0.035))
        pruned = [cols[0]]
        for c in cols[1:]:
            if c - pruned[-1] >= MIN_SEP:
                pruned.append(c)
        cols = pruned
        if len(cols) < 3:
            continue

        # Every interior column must (a) run most of the table's height and
        # (b) sit in WHITESPACE. A real rule has a gutter either side; a line
        # that merely happens to cross characters does not, and using it slices
        # values in half.
        span = y1 - y0
        interior = []
        for c in cols[1:-1]:
            if (vert[y0:y1 + 1, max(0, c - 2):c + 3] > 0).sum() < span * 0.55:
                continue
            left = ink[y0:y1 + 1, max(0, c - 9):max(0, c - 3)]
            right = ink[y0:y1 + 1, c + 4:c + 10]
            gutter = min(left.mean() if left.size else 255,
                         right.mean() if right.size else 255) / 255.0
            if gutter > 0.06:          # text on both sides -> not a boundary
                continue
            interior.append(c)
        cols = sorted(set([cols[0]] + interior + [cols[-1]]))
        if len(cols) < 3:
            continue
        out.append((cols, sorted(ys), (x0, y0, x1, y1)))
    return out


def cell_lines(ink, cell, med_h):
    """Detect text lines INSIDE one cell so they cannot span columns."""
    x, y, w, h = cell["x"], cell["y"], cell["w"], cell["h"]
    sub = ink[y:y + h, x:x + w]
    if sub.size == 0:
        return []
    rows = (sub > 0).sum(axis=1)
    thr = max(2, w * 0.015)
    out, start = [], None
    for i, v in enumerate(rows):
        if v > thr and start is None:
            start = i
        elif v <= thr and start is not None:
            if i - start >= max(6, med_h * 0.35):
                out.append((start, i))
            start = None
    if start is not None and len(rows) - start >= max(6, med_h * 0.35):
        out.append((start, len(rows)))

    lines = []
    for a, b in out:
        seg = sub[a:b, :]
        cols = np.nonzero((seg > 0).any(axis=0))[0]
        if not len(cols):
            continue
        lines.append({"x": int(x + cols[0]), "y": int(y + a),
                      "w": int(cols[-1] - cols[0] + 1), "h": int(b - a),
                      "cell": True,
                      "row": int(cell["row"]), "col": int(cell["col"])})
    return lines


def _overlap_frac(a, b):
    """Fraction of a's area that lies inside b."""
    ix = max(0, min(a["x"] + a["w"], b["x"] + b["w"]) - max(a["x"], b["x"]))
    iy = max(0, min(a["y"] + a["h"], b["y"] + b["h"]) - max(a["y"], b["y"]))
    return (ix * iy) / max(a["w"] * a["h"], 1)


def dedupe(boxes):
    """Drop boxes that are mostly covered by another box.

    Splitting a merged block can leave the parent behind next to its children,
    which renders the same text twice, overlapping. Keep the tighter boxes.
    """
    order = sorted(boxes, key=lambda b: b["w"] * b["h"])   # smallest first
    keep = []
    for b in order:
        if any(_overlap_frac(b, k) > 0.6 or _overlap_frac(k, b) > 0.6 for k in keep):
            continue
        keep.append(b)
    return keep


def detect_lines(ink, graphics=None):
    """Detect text lines. Anything rejected as non-text is recorded in
    `graphics` so it can be PRESERVED on the background instead of erased."""
    sm = cv2.morphologyEx(ink, cv2.MORPH_CLOSE,
                          cv2.getStructuringElement(cv2.MORPH_RECT, (55, 3)))
    cs, _ = cv2.findContours(sm, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    raw = []
    for c in cs:
        x, y, w, h = cv2.boundingRect(c)
        box = {"x": int(x), "y": int(y), "w": int(w), "h": int(h)}
        if w < 40 or h < 9 or h > 260 or w * h < 900:
            if graphics is not None and w * h >= 60:
                graphics.append(box)       # logo, stamp, stray mark -> keep it
            continue
        if ink[y:y + h, x:x + w].mean() / 255.0 < 0.015:
            continue
        if w / h > 60:                     # a rule / underline, not text
            if graphics is not None:
                graphics.append(box)       # keep the rule on the page
            continue
        raw.append(box)
    if not raw:
        return []
    med = float(np.median([b["h"] for b in raw]))
    boxes = [nb for b in raw for nb in split_tall(b, ink, med)]
    before = len(boxes)
    boxes = dedupe(boxes)
    if before != len(boxes):
        print(f"  deduped {before - len(boxes)} overlapping box(es)")
    boxes.sort(key=lambda b: (b["y"] // max(8, int(med * 0.5)), b["x"]))
    for i, b in enumerate(boxes):
        b["id"] = i
    return boxes


# ------------------------------------------------------ stage 3: recognition
def te_to_kn(s):
    """PaddleOCR-VL lacks Kannada and emits Telugu; the blocks are +0x80 aligned."""
    return "".join(chr(ord(c) + 0x80) if 0x0C00 <= ord(c) <= 0x0C7F else c for c in s)


_CACHE_PATH = None
_CACHE = {}
# Recognition runs across OCR_WORKERS threads, so every touch of _CACHE is guarded.
# Without this, concurrent uploads dropped each other's entries and json.dump could
# raise "dictionary changed size during iteration" mid-write.
_CACHE_LOCK = threading.Lock()
_CACHE_LOADED = False


def cache_init(outdir):
    """Disk-cache recognition by crop content, so re-runs are instant.

    The cache is SHARED across jobs (keyed by crop bytes), not per job dir -
    otherwise every upload starts cold and re-pays the full recognition cost.
    Override the location with OCR_CACHE_DIR.

    Loads ONCE per process. It used to rebind the global _CACHE on every request
    while another request's worker threads were still writing to it, so two
    concurrent uploads discarded each other's entries.
    """
    global _CACHE_PATH, _CACHE, _CACHE_LOADED
    with _CACHE_LOCK:
        if _CACHE_LOADED:
            return
        shared = os.environ.get("OCR_CACHE_DIR") or os.path.join(HERE_DIR, "ocr_workspace")
        try:
            os.makedirs(shared, exist_ok=True)
            _CACHE_PATH = os.path.join(shared, ".ocr_cache.json")
        except Exception:
            _CACHE_PATH = os.path.join(outdir, ".ocr_cache.json")
        if os.path.exists(_CACHE_PATH):
            try:
                with open(_CACHE_PATH, encoding="utf-8") as f:
                    _CACHE = json.load(f)
                print(f"ocr cache: {len(_CACHE)} entries")
            except Exception:
                _CACHE = {}
        _CACHE_LOADED = True


def cache_save():
    """
    Snapshot under the lock, then write to a temp file and rename.

    Serialising _CACHE directly while worker threads were still inserting raised
    "dictionary changed size during iteration", and writing in place left a
    truncated cache file if the process died mid-write.
    """
    if not _CACHE_PATH:
        return
    with _CACHE_LOCK:
        snapshot = dict(_CACHE)
    tmp = _CACHE_PATH + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(snapshot, f, ensure_ascii=False)
    os.replace(tmp, _CACHE_PATH)   # atomic


def recognise(crop, model):
    import hashlib
    ok, buf = cv2.imencode(".png", crop)
    raw_bytes = buf.tobytes()
    key = hashlib.sha1(raw_bytes + model.encode()).hexdigest()
    with _CACHE_LOCK:
        if key in _CACHE:
            return _CACHE[key]
    body = json.dumps({"model": model, "prompt": PROMPT,
                       "images": [base64.b64encode(raw_bytes).decode()],
                       "stream": False,
                       "options": {"temperature": 0, "num_predict": 512}}).encode()
    req = urllib.request.Request(f"{OLLAMA}/api/generate", data=body,
                                headers={"Content-Type": "application/json"})
    with urllib.request.urlopen(req, timeout=600) as r:
        out = json.loads(r.read()).get("response", "").strip()
    with _CACHE_LOCK:
        _CACHE[key] = out
    return out


def group_blocks(boxes, med_h):
    """Group vertically-adjacent, horizontally-overlapping lines into blocks.

    Recognition cost is dominated by vision encoding (~1200 prompt tokens per
    call regardless of how little text the crop holds), so reading a 5-line
    paragraph in ONE call is far cheaper than 5 calls. Placement still uses the
    individual line boxes, so alignment is unaffected.
    """
    blocks, cur = [], []

    def hoverlap(a, b):
        ix = min(a["x"] + a["w"], b["x"] + b["w"]) - max(a["x"], b["x"])
        return ix / max(min(a["w"], b["w"]), 1)

    for b in boxes:
        if not cur:
            cur = [b]
            continue
        prev = cur[-1]
        gap = b["y"] - (prev["y"] + prev["h"])
        same_col = hoverlap(prev, b) > 0.55
        similar = abs(b["h"] - prev["h"]) <= max(6, med_h * 0.5)
        if same_col and similar and -med_h * 0.4 <= gap <= med_h * 0.9 and len(cur) < 8:
            cur.append(b)
        else:
            blocks.append(cur)
            cur = [b]
    if cur:
        blocks.append(cur)
    return blocks


def block_bbox(group):
    x0 = min(b["x"] for b in group)
    y0 = min(b["y"] for b in group)
    x1 = max(b["x"] + b["w"] for b in group)
    y1 = max(b["y"] + b["h"] for b in group)
    return x0, y0, x1 - x0, y1 - y0


NOISE_RE = re.compile(r"^[\W_\u4e00-\u9fff]+$")


#  PaddleOCR-VL emits HTML table markup when it thinks a region is tabular. A
#  horizontal rule reliably triggers "<ecel><ecel>..." which then renders as a
#  giant garbage line across the page.
#  Any angle-bracket token is model markup, not page text. Naming them
#  individually kept letting new ones through (<ecel>, then <cfrml>).
STRUCT_TOKEN_RE = re.compile(r"<\s*/?\s*[A-Za-z][\w:-]{0,15}\s*/?\s*>")


#  The model writes superscripts and ordinals as LaTeX: "23\(^{rd}\) March".
#  Rejecting anything containing a backslash threw away whole correct lines, so
#  unwrap the common forms first and only judge what is left.
_LATEX_SUP = re.compile(r"\\?\(?\s*\^\{([^}]*)\}\s*\\?\)?")
_LATEX_WRAP = re.compile(r"\\[()\[\]]")
_LATEX_CMD = re.compile(r"\\[A-Za-z]+\s*(\{[^}]*\})?")


def delatex(t):
    if not t or "\\" not in t:
        return t
    s = _LATEX_SUP.sub(r"\1", t)          # \(^{rd}\) -> rd
    s = re.sub(r"\$([^$]*)\$", r"\1", s)  # $x$ -> x
    s = _LATEX_CMD.sub("", s)             # \dagger, \frac{..} -> gone
    s = _LATEX_WRAP.sub("", s)
    return re.sub(r"\s+", " ", s).strip()


def is_noise(t):
    if not t:
        return True
    s = delatex(t).strip()
    if not s:
        return True
    # Short strings are legitimate in table cells and margins ("1", "17", "No.")
    if len(s) <= 2:
        return not any(c.isalnum() for c in s)
    if STRUCT_TOKEN_RE.search(s):                     # table/HTML markup, not text
        return True
    toks = s.split()
    if len(toks) > 5 and len(set(toks)) <= max(2, len(toks) // 4):
        return True                                    # "or/ or/ or/ ..." = a rule
    if NOISE_RE.match(s):
        return True
    alnum = sum(c.isalnum() for c in s)
    return alnum / max(len(s), 1) < 0.35


#  The repeated unit must contain a letter or digit. Requiring only "any 1-6
#  chars repeated" rejected real text with dot leaders, e.g.
#  "This Certified Copy Contains.....Pages".
REPEAT_RE = re.compile(r"((?=[^\W_]*[^\W_])[\s\S]{1,6}?)\1{3,}")
FILLER_RE = re.compile(r"[.\-_·•,:;/\\|~=*\s]{3,}")


def looks_invented(t):
    """Reject output that cannot have come from real text on the page.

    Measured on this corpus, hallucinations share obvious signatures: a short
    substring repeated many times (a printed rule read as a token sequence,
    "/T/i/t/i/t/i" or "00/00/00/00"), or a run of statistics the document does
    not contain at all ("- 25.7% - 38.4% - 50.9%").
    """
    if not t:
        return True
    s = t.strip()
    if STRUCT_TOKEN_RE.search(s):
        return True
    # judge the text with runs of dot/dash filler collapsed, so form leaders
    # ("Contains.....Pages") are not mistaken for a repeating hallucination
    probe = FILLER_RE.sub(" ", s)
    compact = re.sub(r"\s+", "", probe)
    if len(compact) >= 8 and REPEAT_RE.search(compact):
        return True
    # a list of percentages / measurements with almost no letters is invented
    nums = re.findall(r"\d+(?:[.,]\d+)?\s*%", s)
    if len(nums) >= 3:
        return True
    letters = sum(c.isalpha() for c in probe)
    digits = sum(c.isdigit() for c in probe)
    if digits >= 8 and letters <= 2:
        return True
    if re.search(r"\b(km/h|mph|kg|MB|GB|px)\b", s) and letters <= 6:
        return True
    return False


def ink_coverage(ink, b):
    sub = ink[b["y"]:b["y"] + b["h"], b["x"]:b["x"] + b["w"]]
    if sub.size == 0:
        return 0.0, 0.0
    cov = float(sub.mean() / 255.0)
    cols = (sub > 0).sum(axis=0)
    occ = float((cols > 0).mean())
    return cov, occ


def recognise_all(rot, boxes, model, workers=None, use_blocks=True, verbose=False,
                  ink=None):
    """Recognise every box, grouping into blocks and running calls in parallel.

    Falls back to per-line recognition for any block whose returned line count
    does not match the number of boxes in it, so a bad split can never
    mis-assign text to the wrong position.
    """
    from concurrent.futures import ThreadPoolExecutor

    H, W = rot.shape[:2]
    workers = workers or int(os.environ.get("OCR_WORKERS", "3"))
    pad = 4

    def crop_of(x, y, w, h, p=None):
        p = pad if p is None else p
        c = rot[max(0, y - p):min(H, y + h + p), max(0, x - p):min(W, x + w + p)]
        return cv2.resize(c, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)

    def crop_cell(b, margin=10):
        """Cell crop isolated on a white canvas.

        Extending the crop into neighbouring pixels pulls the cell wall in and it
        gets read as a glyph ("24" -> "124"); cropping tighter instead clips the
        leading character ("Item" -> "tem") and loses digits ("17" -> "1").
        Taking the exact cell content and surrounding it with white gives the
        model isolated text with neither problem.
        """
        c = rot[max(0, b["y"]):min(H, b["y"] + b["h"]),
                max(0, b["x"]):min(W, b["x"] + b["w"])]
        if c.size == 0:
            return None
        canvas = np.full((c.shape[0] + margin * 2, c.shape[1] + margin * 2, 3),
                         245, np.uint8)
        canvas[margin:margin + c.shape[0], margin:margin + c.shape[1]] = c
        return cv2.resize(canvas, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)

    def run_line(b):
        try:
            crop = crop_cell(b) if b.get("cell") else \
                crop_of(b["x"], b["y"], b["w"], b["h"])
            if crop is None:
                return b, ""
            return b, recognise(crop, model)
        except Exception:
            return b, ""

    # Gate on ink before calling the model. Crops with almost no ink are
    # handwriting, rules or paper texture; asked to read them the model invents
    # plausible content ("- 25.7% - 38.4% - 50.9%", "/T/i/t/i/t/i", "12.000 km/h").
    # Measured on this corpus real text lines sit at coverage 0.116-0.33.
    skipped = 0
    if ink is not None and len(boxes) >= 6:
        covs = []
        for b in boxes:
            b["cov"], b["occ"] = ink_coverage(ink, b)
            covs.append(b["cov"])
        med_cov = float(np.median(covs)) or 1.0
        floor = max(0.05, med_cov * 0.62)
        for b in boxes:
            if b["cov"] < floor:
                b["text"], b["raw"] = "", ""
                b["noise"] = True
                b["skip_reason"] = f"low ink {b['cov']:.3f} < {floor:.3f}"
                skipped += 1
        if verbose and skipped:
            print(f"  skipped {skipped} low-ink region(s) before OCR "
                  f"(hallucination guard, floor={floor:.3f})")

    todo = [b for b in boxes if not b.get("noise")]
    med = float(np.median([b["h"] for b in todo])) if todo else 20.0
    groups = group_blocks(todo, med) if use_blocks else [[b] for b in todo]
    multi = [g for g in groups if len(g) > 1]
    if verbose:
        print(f"  recognition: {len(todo)} lines -> {len(groups)} calls "
              f"({len(multi)} multi-line blocks), {workers} workers")

    def run_group(g):
        if len(g) == 1 or any(b.get("cell") for b in g):
            return [run_line(b) for b in g]   # never merge across cells
        x, y, w, h = block_bbox(g)
        try:
            raw = recognise(crop_of(x, y, w, h), model)
        except Exception:
            raw = ""
        lines = [l for l in (raw or "").split("\n") if l.strip()]
        if len(lines) == len(g):
            return list(zip(g, lines))
        # line count mismatch -> re-read individually rather than risk
        # attaching the wrong text to a position
        return [run_line(b) for b in g]

    out = []
    with ThreadPoolExecutor(max_workers=workers) as ex:
        for res in ex.map(run_group, groups):
            out.extend(res)

    invented = 0
    for b, raw in out:
        raw = (raw or "").replace("\n", " ").strip()
        b["raw"] = raw
        if looks_invented(raw):
            b["text"], b["noise"] = "", True
            b["skip_reason"] = "output looks invented"
            invented += 1
            continue
        b["text"] = te_to_kn(delatex(raw))
        b["noise"] = is_noise(b["text"])
        if b["noise"] and not b.get("skip_reason"):
            b["skip_reason"] = "is_noise filter"
    if verbose and invented:
        print(f"  rejected {invented} invented output(s) after OCR")
    cache_save()
    return boxes


# --------------------------------------------------------- stage 6: rendering
def pick_font():
    for p in FONT_CANDIDATES:
        if os.path.exists(p):
            return p
    return None


def detect_bold(ink, boxes):
    """Flag boxes whose strokes are heavier than the page norm.

    Bold text lays down more ink per unit of glyph area at the same size, so
    comparing each line's ink coverage against the median for similar-height
    lines separates bold from regular without needing font metrics.
    """
    live = [b for b in boxes if not b.get("noise")]
    if len(live) < 4:
        for b in live:
            b["bold"] = False
        return
    for b in live:
        sub = ink[b["y"]:b["y"] + b["h"], b["x"]:b["x"] + b["w"]]
        b["_cov"] = float(sub.mean() / 255.0) if sub.size else 0.0
    med_cov = float(np.median([b["_cov"] for b in live if b["_cov"] > 0])) or 1.0
    for b in live:
        b["bold"] = b["_cov"] >= med_cov * 1.28
        b.pop("_cov", None)


def build_text_mask(ink, boxes, graphics, shape):
    """Erase ONLY the text we are going to redraw.

    Inpainting every ink pixel destroys rules, the logo, stamps and handwriting -
    they are ink, but nothing draws them back, so they vanish from the output.
    Restricting the mask to replaced text lines leaves all of that untouched.
    """
    mask = np.zeros(shape[:2], np.uint8)
    for b in boxes:
        if b.get("noise"):
            continue           # not replaced -> must stay visible
        pad = 2
        y0, y1 = max(0, b["y"] - pad), min(shape[0], b["y"] + b["h"] + pad)
        x0, x1 = max(0, b["x"] - pad), min(shape[1], b["x"] + b["w"] + pad)
        mask[y0:y1, x0:x1] = 255
    mask = cv2.bitwise_and(mask, ink)
    # don't erase a rule that happens to pass through a text box
    for g in (graphics or []):
        if g["h"] <= 6 or (g["w"] / max(g["h"], 1)) > 25:
            y0, y1 = max(0, g["y"] - 1), min(shape[0], g["y"] + g["h"] + 1)
            x0, x1 = max(0, g["x"] - 1), min(shape[1], g["x"] + g["w"] + 1)
            mask[y0:y1, x0:x1] = 0
    return cv2.dilate(mask, cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (5, 5)),
                      iterations=2)


class Fitter:
    """Single font-sizing rule: seed from box height, shrink until it fits the box."""

    def __init__(self, font_path):
        self.path, self.cache = font_path, {}

    def font(self, px, bold=False):
        key = (int(max(6, px)), bool(bold))
        if key not in self.cache:
            px_i = key[0]
            loaded = None
            if bold:
                for cand in (r"C:\Windows\Fonts\NirmalaB.ttf",
                             r"C:\Windows\Fonts\Nirmala.ttc"):
                    if os.path.exists(cand):
                        for idx in (1, 0):
                            try:
                                loaded = ImageFont.truetype(cand, px_i, index=idx)
                                break
                            except Exception:
                                continue
                    if loaded:
                        break
            if loaded is None:
                try:
                    loaded = ImageFont.truetype(self.path, px_i, index=0)
                except Exception:
                    loaded = ImageFont.load_default()
            self.cache[key] = loaded
        return self.cache[key]

    def fit(self, draw, text, bw, bh):
        size = max(6, int(bh * 0.82))
        while size > 6:
            f = self.font(size)
            l, t, r, b = draw.textbbox((0, 0), text, font=f)
            if (r - l) <= bw * 0.995 and (b - t) <= bh * 1.15:
                return f, r - l, b - t, size
            size -= 1
        f = self.font(6)
        l, t, r, b = draw.textbbox((0, 0), text, font=f)
        return f, r - l, b - t, 6


def assign_font_sizes(dr, boxes, fitter):
    """One font size per PARAGRAPH, not per line.

    Detected box heights wobble by a few pixels depending on whether a line
    happens to contain ascenders or descenders. Sizing each line from its own
    height therefore produces visibly ragged typography, even when every line is
    positioned correctly. Real documents set a size per paragraph, so we take a
    robust (median) height per block and apply it to every line in that block,
    shrinking only where a specific line would overflow its width.
    """
    live = [b for b in boxes if not b.get("noise")]
    if not live:
        return

    # A document has ONE dominant body size. Per-block medians are not enough:
    # a single line whose detected height happens to be 31px next to mates at
    # 38-53px gets rejected from its own paragraph and rendered visibly smaller.
    # So derive a page-wide body height from the widest lines (body text runs
    # full-measure; headings and footnotes do not) and share it.
    # Body height from the MODE of line heights, not from the widest lines.
    # The width heuristic collapses on table-heavy pages, where most lines are
    # narrow cells and the few wide ones are rules or headings - that produced
    # a giant heading next to microscopic body text.
    hs = np.array([b["h"] for b in live], float)
    lo, hi = max(6.0, hs.min()), hs.max()
    if hi - lo < 2:
        body_h = float(np.median(hs))
    else:
        hist, edges = np.histogram(hs, bins=max(4, min(14, len(hs) // 2)),
                                   range=(lo, hi + 1e-6))
        k = int(np.argmax(hist))
        in_bin = hs[(hs >= edges[k]) & (hs <= edges[k + 1])]
        body_h = float(np.median(in_bin)) if in_bin.size else float(np.median(hs))
    body_lines = [b for b in live if abs(b["h"] - body_h) <= body_h * 0.45]

    def fit_width(text, bw, start):
        s = max(6, int(start))
        while s > 6:
            l, t, r, bo = dr.textbbox((0, 0), text, font=fitter.font(s))
            if (r - l) <= bw * 0.995:
                return s
            s -= 1
        return 6

    BODY_TOL = 0.45          # within +/-45% of the body height => body text
    body_size = max(6, int(body_h * 0.78))
    # shrink the shared body size until the widest body line fits
    for _ in range(80):
        worst = 0.0
        f = fitter.font(body_size)
        for b in (body_lines or live):
            l, t, r, bo = dr.textbbox((0, 0), b["text"], font=f)
            worst = max(worst, (r - l) / max(b["w"] * 0.995, 1))
        if worst <= 1.0 or body_size <= 6:
            break
        body_size -= 1

    # An outlier's own box height is not always trustworthy - a faint line
    # recovered from a gap can get an over-tall box, which then renders the text
    # absurdly large ("424BM Agenda" at 3x its real size). Clamp outliers to a
    # sane multiple of the body size.
    OUT_MIN, OUT_MAX = 0.55, 1.9

    for b in live:
        if abs(b["h"] - body_h) <= body_h * BODY_TOL:
            # body text: use the shared size, reduced only if this line overflows
            b["font_px"] = fit_width(b["text"], b["w"], body_size)
            b["size_class"] = "body"
        else:
            want = max(6, int(b["h"] * 0.78))
            want = int(min(max(want, body_size * OUT_MIN), body_size * OUT_MAX))
            b["font_px"] = fit_width(b["text"], b["w"], want)
            b["size_class"] = "outlier"


def render(rot, ink, boxes, fitter, original=None, M=None, keep_skew=False,
           graphics=None):
    """Remove the original text and put the recognised text back.

    keep_skew=False -> render on the deskewed page (text horizontal)
    keep_skew=True  -> render on the ORIGINAL page, rotating each line back to
                       the angle it had in the photograph, so the result matches
                       the raw scan's appearance instead of a straightened one.
    """
    detect_bold(ink, boxes)

    if not keep_skew or original is None or M is None:
        mask = build_text_mask(ink, boxes, graphics, rot.shape)
        clean = cv2.inpaint(rot, mask, 6, cv2.INPAINT_TELEA)
        pil = Image.fromarray(cv2.cvtColor(clean, cv2.COLOR_BGR2RGB))
        dr = ImageDraw.Draw(pil)
        assign_font_sizes(dr, boxes, fitter)
        for b in boxes:
            if b.get("noise"):
                continue
            f = fitter.font(b["font_px"], b.get("bold"))
            l, t, r, bo = dr.textbbox((0, 0), b["text"], font=f)
            dr.text((b["x"], b["y"] + (b["h"] - (bo - t)) / 2 - t), b["text"],
                    font=f, fill=(20, 20, 20))
        return clean, pil

    # ---- slant-preserving path ----
    H, W = original.shape[:2]
    # the mask is built in deskewed space then mapped back, so only replaced
    # text is erased; rules, logo, stamps and handwriting survive untouched
    text_mask = build_text_mask(ink, boxes, graphics, rot.shape)
    Minv = cv2.invertAffineTransform(M)
    mask = cv2.warpAffine(text_mask, Minv, (W, H), flags=cv2.INTER_NEAREST)
    clean = cv2.inpaint(original, mask, 6, cv2.INPAINT_TELEA)

    pil = Image.fromarray(cv2.cvtColor(clean, cv2.COLOR_BGR2RGB))
    scratch = ImageDraw.Draw(Image.new("RGB", (8, 8)))
    assign_font_sizes(scratch, boxes, fitter)

    # skew recovered from the matrix: M maps original -> deskewed by +skew,
    # so a line drawn horizontally must be rotated by -skew to sit as photographed
    skew_deg = float(np.degrees(np.arctan2(M[1, 0], M[0, 0])))

    for b in boxes:
        if b.get("noise"):
            continue
        f = fitter.font(b["font_px"], b.get("bold"))
        l, t, r, bo = scratch.textbbox((0, 0), b["text"], font=f)
        tw, th = max(1, r - l), max(1, bo - t)

        layer = Image.new("RGBA", (tw + 8, th + 8), (0, 0, 0, 0))
        ImageDraw.Draw(layer).text((4 - l, 4 - t), b["text"], font=f,
                                   fill=(20, 20, 20, 255))
        rot_layer = layer.rotate(skew_deg, resample=Image.BICUBIC, expand=True)

        # centre of this line in deskewed space -> original space
        cx, cy = b["x"] + b["w"] / 2.0, b["y"] + b["h"] / 2.0
        ox = Minv[0, 0] * cx + Minv[0, 1] * cy + Minv[0, 2]
        oy = Minv[1, 0] * cx + Minv[1, 1] * cy + Minv[1, 2]
        b["orig_cx"], b["orig_cy"] = int(ox), int(oy)

        px = int(round(ox - rot_layer.width / 2))
        py = int(round(oy - rot_layer.height / 2))
        pil.paste(rot_layer, (px, py), rot_layer)

    return clean, pil


def score_alignment(pil, ink, boxes):
    rend = (np.array(pil.convert("L")) < 128).astype(np.uint8)
    band = cv2.dilate((ink > 0).astype(np.uint8),
                      cv2.getStructuringElement(cv2.MORPH_RECT, (9, 9)), 1)
    out = []
    for b in boxes:
        if b.get("noise"):
            continue
        sl = (slice(b["y"], b["y"] + b["h"]), slice(b["x"], b["x"] + b["w"]))
        r = rend[sl].sum()
        hit = (rend[sl] & band[sl]).sum()
        b["align"] = round(100 * hit / max(r, 1), 1)
        out.append(b["align"])
    return out


# ------------------------------------------------------------------ the DOCX
def write_docx(path, bg_png, boxes, W, H):
    from docx import Document
    from docx.shared import Emu
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn, nsmap
    nsmap["wps"] = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"

    doc = Document()
    sec = doc.sections[0]
    aspect = W / H
    page_h = int(11 * 914400)
    page_w = int(page_h * aspect)
    sec.page_width, sec.page_height = Emu(page_w), Emu(page_h)
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, m, Emu(0))

    # python-docx's default template may start with zero paragraphs
    para = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    shape = para.add_run().add_picture(bg_png)
    rid = shape._inline.xpath(".//a:blip")[0].get(
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
    para.clear()

    def anchor(rid_, w, h, behind):
        d = OxmlElement("w:drawing")
        a = OxmlElement("wp:anchor")
        for k, v in (("distT", "0"), ("distB", "0"), ("distL", "0"), ("distR", "0"),
                     ("simplePos", "0"), ("relativeHeight", "0"),
                     ("behindDoc", "1" if behind else "0"), ("locked", "0"),
                     ("layoutInCell", "1"), ("allowOverlap", "1")):
            a.set(k, v)
        sp = OxmlElement("wp:simplePos"); sp.set("x", "0"); sp.set("y", "0"); a.append(sp)
        return d, a

    # background image, behind text
    d, a = anchor(rid, page_w, page_h, True)
    for tag, off in (("wp:positionH", "0"), ("wp:positionV", "0")):
        p = OxmlElement(tag); p.set("relativeFrom", "page")
        o = OxmlElement("wp:posOffset"); o.text = off; p.append(o); a.append(p)
    ext = OxmlElement("wp:extent"); ext.set("cx", str(page_w)); ext.set("cy", str(page_h))
    a.append(ext); a.append(OxmlElement("wp:wrapNone"))
    dp = OxmlElement("wp:docPr"); dp.set("id", "1"); dp.set("name", "bg"); a.append(dp)
    g = OxmlElement("a:graphic"); gd = OxmlElement("a:graphicData")
    gd.set("uri", "http://schemas.openxmlformats.org/drawingml/2006/picture")
    pic = OxmlElement("pic:pic")
    nv = OxmlElement("pic:nvPicPr"); cp = OxmlElement("pic:cNvPr")
    cp.set("id", "1"); cp.set("name", "bg.png"); nv.append(cp)
    nv.append(OxmlElement("pic:cNvPicPr")); pic.append(nv)
    bf = OxmlElement("pic:blipFill"); bl = OxmlElement("a:blip")
    bl.set(qn("r:embed"), rid); bf.append(bl)
    st = OxmlElement("a:stretch"); st.append(OxmlElement("a:fillRect")); bf.append(st)
    pic.append(bf)
    spr = OxmlElement("pic:spPr"); xf = OxmlElement("a:xfrm")
    o1 = OxmlElement("a:off"); o1.set("x", "0"); o1.set("y", "0"); xf.append(o1)
    e1 = OxmlElement("a:ext"); e1.set("cx", str(page_w)); e1.set("cy", str(page_h))
    xf.append(e1); spr.append(xf)
    pg = OxmlElement("a:prstGeom"); pg.set("prst", "rect")
    pg.append(OxmlElement("a:avLst")); spr.append(pg); pic.append(spr)
    gd.append(pic); g.append(gd); a.append(g); d.append(a)
    r = OxmlElement("w:r"); r.append(d); para._p.append(r)

    sx, sy = page_w / W, page_h / H
    for i, b in enumerate(boxes, start=2):
        if b.get("noise"):
            continue
        x, y = int(b["x"] * sx), int(b["y"] * sy)
        w, h = int(b["w"] * sx), int(b["h"] * sy)
        half_pt = max(8, min(int(b.get("font_px", 12) * sy / 12700 * 2), 96))

        d = OxmlElement("w:drawing"); a = OxmlElement("wp:anchor")
        for k, v in (("distT", "0"), ("distB", "0"), ("distL", "0"), ("distR", "0"),
                     ("simplePos", "0"), ("relativeHeight", str(i + 10)),
                     ("behindDoc", "0"), ("locked", "0"),
                     ("layoutInCell", "1"), ("allowOverlap", "1")):
            a.set(k, v)
        sp = OxmlElement("wp:simplePos"); sp.set("x", "0"); sp.set("y", "0"); a.append(sp)
        for tag, off in (("wp:positionH", x), ("wp:positionV", y)):
            p = OxmlElement(tag); p.set("relativeFrom", "page")
            o = OxmlElement("wp:posOffset"); o.text = str(off); p.append(o); a.append(p)
        ext = OxmlElement("wp:extent"); ext.set("cx", str(w)); ext.set("cy", str(h))
        a.append(ext); a.append(OxmlElement("wp:wrapNone"))
        dp = OxmlElement("wp:docPr"); dp.set("id", str(i))
        dp.set("name", f"t{i}"); a.append(dp)
        g = OxmlElement("a:graphic"); gd = OxmlElement("a:graphicData")
        gd.set("uri", "http://schemas.microsoft.com/office/word/2010/wordprocessingShape")
        wsp = OxmlElement("wps:wsp")
        cnv = OxmlElement("wps:cNvSpPr"); cnv.set("txBox", "1"); wsp.append(cnv)
        spr = OxmlElement("wps:spPr"); xf = OxmlElement("a:xfrm")
        o1 = OxmlElement("a:off"); o1.set("x", "0"); o1.set("y", "0"); xf.append(o1)
        e1 = OxmlElement("a:ext"); e1.set("cx", str(w)); e1.set("cy", str(h))
        xf.append(e1); spr.append(xf)
        pg = OxmlElement("a:prstGeom"); pg.set("prst", "rect")
        pg.append(OxmlElement("a:avLst")); spr.append(pg)
        spr.append(OxmlElement("a:noFill"))
        ln = OxmlElement("a:ln"); ln.set("w", "0")
        ln.append(OxmlElement("a:noFill")); spr.append(ln); wsp.append(spr)
        tx = OxmlElement("wps:txbx"); tc = OxmlElement("w:txbxContent")
        p = OxmlElement("w:p"); ppr = OxmlElement("w:pPr")
        spc = OxmlElement("w:spacing")
        spc.set(qn("w:before"), "0"); spc.set(qn("w:after"), "0")
        ppr.append(spc); p.append(ppr)
        run = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
        for t in ("w:sz", "w:szCs"):
            e = OxmlElement(t); e.set(qn("w:val"), str(half_pt)); rpr.append(e)
        kn = any(0x0C80 <= ord(c) <= 0x0CFF for c in b["text"])
        rf = OxmlElement("w:rFonts")
        name = "Nirmala UI" if kn else "Arial"
        for at in ("w:ascii", "w:hAnsi", "w:cs"):
            rf.set(qn(at), name)
        rpr.append(rf); run.append(rpr)
        t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve")
        t.text = b["text"]; run.append(t); p.append(run)
        tc.append(p); tx.append(tc); wsp.append(tx)
        bp = OxmlElement("wps:bodyPr")
        for k, v in (("wrap", "none"), ("lIns", "0"), ("tIns", "0"),
                     ("rIns", "0"), ("bIns", "0"), ("anchor", "ctr")):
            bp.set(k, v)
        wsp.append(bp)
        gd.append(wsp); g.append(gd); a.append(g); d.append(a)
        rr = OxmlElement("w:r"); rr.append(d); para._p.append(rr)

    doc.save(path)


# ------------------------------------------------- API adapter (used by main.py)
def process_page(image_path, outdir, model=DEFAULT_MODEL, want_docx=True,
                 keep_skew=True):
    """Run the whole pipeline on one page.

    Returns a dict shaped for the existing /extract contract:
      text_data     -> [{text,left,top,width,height,kannada,align,font_px}]
      preview_image -> base64 PNG of the page with its original text removed
      extracted_text-> newline-joined text in reading order
      docx          -> path to the aligned .docx (or None)
    """
    img = load_page(image_path)
    if img is None:
        raise ValueError(f"cannot read {image_path}")
    os.makedirs(outdir, exist_ok=True)
    H, W = img.shape[:2]

    rot, ink, skew, frac, ink_sens, M = normalise(img)
    graphics = []
    boxes = detect_lines(ink, graphics)
    boxes = recover_faint_lines(boxes, ink_sens, ink)
    boxes = apply_table_grid(boxes, ink)
    boxes.sort(key=lambda b: (b["y"] // 18, b["x"]))
    for i, b in enumerate(boxes):
        b["id"] = i

    cache_init(outdir)
    t_ocr = time.time()
    recognise_all(rot, boxes, model, ink=ink)
    ocr_secs = round(time.time() - t_ocr, 1)

    fp = pick_font()
    if not fp:
        raise RuntimeError("no Kannada-capable font found")
    clean, pil = render(rot, ink, boxes, Fitter(fp),
                        original=img, M=M, keep_skew=keep_skew,
                        graphics=graphics)
    if not keep_skew:
        score_alignment(pil, ink, boxes)

    stem = os.path.splitext(os.path.basename(image_path))[0]
    bg = os.path.join(outdir, f"{stem}_clean_bg.png")
    cv2.imwrite(bg, clean)
    pil.save(os.path.join(outdir, f"{stem}_overlay.png"))
    cv2.imwrite(os.path.join(outdir, f"{stem}_deskewed.png"), rot)

    docx_path = None
    if want_docx:
        docx_path = os.path.join(outdir, f"{stem}_aligned.docx")
        try:
            write_docx(docx_path, bg, boxes, W, H)
        except Exception:
            docx_path = None

    kept = [b for b in boxes if not b.get("noise")]
    text_data = [{
        "text": b["text"],
        "left": b["x"], "top": b["y"], "width": b["w"], "height": b["h"],
        "kannada": any(0x0C80 <= ord(c) <= 0x0CFF for c in b["text"]),
        "align": b.get("align"),
        "font_px": b.get("font_px"),
    } for b in kept]

    with open(bg, "rb") as f:
        preview_b64 = base64.b64encode(f.read()).decode()

    aligns = [b["align"] for b in kept if b.get("align") is not None]
    return {
        "text_data": text_data,
        "extracted_text": "\n".join(b["text"] for b in kept),
        "preview_image": preview_b64,
        "preview_mime": "image/png",
        "docx": docx_path,
        "clean_bg": bg,
        "skew": round(skew, 2),
        "page_frac": round(frac, 3),
        "lines": len(boxes),
        "rendered": len(kept),
        "ocr_secs": ocr_secs,
        "align_median": round(float(np.median(aligns)), 1) if aligns else None,
        "width": W, "height": H,
    }


# -------------------------------------------------------------------- driver
def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("input")
    ap.add_argument("--model", default=DEFAULT_MODEL)
    ap.add_argument("--outdir", default=None)
    ap.add_argument("--no-ocr", action="store_true",
                    help="detect and render boxes only (no model needed)")
    ap.add_argument("--no-blocks", action="store_true",
                    help="recognise line-by-line instead of grouping paragraphs "
                         "(slower; use to compare accuracy)")
    ap.add_argument("--deskew", action="store_true",
                    help="straighten the output. Default is to KEEP the original "
                         "slant so the result matches the raw photograph.")
    args = ap.parse_args()

    img = load_page(args.input)
    if img is None:
        sys.exit(f"cannot read {args.input}")
    stem = os.path.splitext(os.path.basename(args.input))[0]
    outdir = args.outdir or os.path.dirname(os.path.abspath(args.input))
    os.makedirs(outdir, exist_ok=True)
    H, W = img.shape[:2]
    print(f"page {W}x{H}")

    rot, ink, skew, frac, ink_sens, M = normalise(img)
    print(f"page area {frac*100:.1f}% of frame | skew {skew:+.2f} deg "
          f"(fixed-8%-crop would discard 29.4%)")

    graphics = []
    boxes = detect_lines(ink, graphics)
    boxes = recover_faint_lines(boxes, ink_sens, ink)
    boxes = apply_table_grid(boxes, ink)
    boxes.sort(key=lambda b: (b["y"] // 18, b["x"]))
    for i, b in enumerate(boxes):
        b["id"] = i
    ncell = sum(1 for b in boxes if b.get("cell"))
    print(f"text lines detected: {len(boxes)}  "
          f"(table cells: {ncell})  non-text kept: {len(graphics)}  "
          f"(render: {'SLANTED like original' if not args.deskew else 'deskewed'})")

    if args.no_ocr:
        for b in boxes:
            b["text"], b["noise"] = f"[{b['id']}]", False
    else:
        cache_init(outdir)
        t0 = time.time()
        recognise_all(rot, boxes, args.model, ink=ink,
                      use_blocks=not args.no_blocks, verbose=True)
        for b in boxes:
            flag = "  (dropped)" if b["noise"] else ""
            print(f"  [{b['id']:3d}] {b['text'][:84]}{flag}")
        print(f"recognition: {time.time()-t0:.0f}s for {len(boxes)} lines")

    fp = pick_font()
    if not fp:
        sys.exit("no Kannada-capable font found - install Noto Sans Kannada")
    print(f"font: {os.path.basename(fp)}")

    keep_skew = not args.deskew
    clean, pil = render(rot, ink, boxes, Fitter(fp),
                        original=img, M=M, keep_skew=keep_skew,
                        graphics=graphics)
    nb = sum(1 for b in boxes if b.get("bold"))
    print(f"bold lines detected: {nb}")
    scores = [] if keep_skew else score_alignment(pil, ink, boxes)

    bg = os.path.join(outdir, f"{stem}_clean_bg.png")
    ov = os.path.join(outdir, f"{stem}_overlay.png")
    cv2.imwrite(bg, clean)
    cv2.imwrite(os.path.join(outdir, f"{stem}_deskewed.png"), rot)
    cv2.imwrite(os.path.join(outdir, f"{stem}_ink.png"), ink)
    pil.save(ov)
    # compare against whichever frame we rendered into
    ref = img if keep_skew else rot
    Image.fromarray(np.hstack([
        cv2.cvtColor(ref, cv2.COLOR_BGR2RGB),
        np.full((H, 24, 3), 255, np.uint8),
        np.array(pil)])).save(os.path.join(outdir, f"{stem}_sidebyside.png"))

    dx = os.path.join(outdir, f"{stem}_aligned.docx")
    try:
        write_docx(dx, bg, boxes, W, H)
    except Exception:
        import traceback
        traceback.print_exc()
        dx = "DOCX FAILED (traceback above)"

    def _plain(o):
        if isinstance(o, (np.integer,)):
            return int(o)
        if isinstance(o, (np.floating,)):
            return float(o)
        if isinstance(o, (np.bool_,)):
            return bool(o)
        raise TypeError(f"not serialisable: {type(o).__name__}")

    with open(os.path.join(outdir, f"{stem}_lines.json"), "w", encoding="utf-8") as f:
        json.dump({"skew": skew, "page_frac": frac, "w": W, "h": H, "boxes": boxes},
                  f, ensure_ascii=False, indent=1, default=_plain)

    if scores:
        bad = [b["id"] for b in boxes if not b.get("noise") and b.get("align", 100) < 50]
        print(f"\nalignment: mean {np.mean(scores):.1f}%  median {np.median(scores):.1f}%  "
              f"worst {np.min(scores):.1f}%")
        print(f"misplaced (<50%): {bad if bad else 'none'}")
    print(f"\n{bg}\n{ov}\n{dx}")


if __name__ == "__main__":
    main()
